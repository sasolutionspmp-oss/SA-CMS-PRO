from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from app.ingest.pdf import extract_pdf_text

try:  # pragma: no cover - optional dependency at runtime
    import fitz  # type: ignore
except Exception:  # noqa: BLE001
    fitz = None  # type: ignore

try:  # pragma: no cover - optional dependency
    import docx  # type: ignore
except Exception:  # noqa: BLE001
    docx = None  # type: ignore

try:  # pragma: no cover - optional dependency
    import openpyxl  # type: ignore
except Exception:  # noqa: BLE001
    openpyxl = None  # type: ignore


MAX_PREVIEW_CHARS = 100_000
CSV_PREVIEW_ROWS = 200
XLSX_PREVIEW_ROWS = 200


@dataclass
class ParseOutcome:
    status: str
    text: str
    metadata: dict[str, Any]
    page_count: int | None
    error: str | None = None

    @property
    def snippet(self) -> str | None:
        data = self.text.strip()
        if not data:
            return None
        normalized = " ".join(data.split())
        return normalized[:500]


def parse_document(path: Path, mime_type: str) -> ParseOutcome:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path, mime_type)
    if suffix == ".docx":
        return _parse_docx(path, mime_type)
    if suffix in {".csv"}:
        return _parse_csv(path, mime_type)
    if suffix in {".txt", ".text"}:
        return _parse_text(path, mime_type)
    if suffix in {".xlsx"}:
        return _parse_xlsx(path, mime_type)
    if suffix in {".dwg", ".dxf"}:
        return _unsupported_cad(path, mime_type)
    return ParseOutcome(
        status="failed",
        text="",
        metadata={},
        page_count=None,
        error=f"Unsupported file type: {suffix}",
    )


def _truncate_text(text: str) -> str:
    if len(text) <= MAX_PREVIEW_CHARS:
        return text
    return text[:MAX_PREVIEW_CHARS] + "\n...\n[truncated]"


def _parse_pdf(path: Path, mime_type: str) -> ParseOutcome:
    text = extract_pdf_text(path)
    metadata: dict[str, Any] = {"mime": mime_type}
    page_count: int | None = None
    if fitz is not None:  # pragma: no branch - optional
        try:
            with fitz.open(path) as doc:  # type: ignore[call-arg]
                page_count = doc.page_count
        except Exception:  # noqa: BLE001 - degrade gracefully
            page_count = None
    return ParseOutcome(
        status="parsed",
        text=_truncate_text(text),
        metadata=metadata,
        page_count=page_count,
        error=None if text.strip() else "No extractable text",
    )


def _parse_docx(path: Path, mime_type: str) -> ParseOutcome:
    if docx is None:
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error="python-docx dependency is not installed",
        )
    try:
        document = docx.Document(path)  # type: ignore[attr-defined]
    except Exception as exc:  # noqa: BLE001
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error=str(exc),
        )
    lines: list[str] = []
    for para in document.paragraphs:
        data = para.text.strip()
        if data:
            lines.append(data)
    table_count = 0
    tables_attr = getattr(document, "tables", [])
    for table_index, table in enumerate(tables_attr, 1):
        table_count += 1
        lines.append(f"# Table {table_index}")
        row_iter = getattr(table, "rows", [])
        for row_index, row in enumerate(row_iter, 1):
            cells: list[str] = []
            row_cells = getattr(row, "cells", [])
            for col_index, cell in enumerate(row_cells, 1):
                value = getattr(cell, "text", "").strip()
                if value:
                    cells.append(f"R{row_index}C{col_index}: {value}")
            if cells:
                lines.append(" | ".join(cells))
        lines.append("")
    text = _truncate_text("\n".join(lines))
    metadata: dict[str, Any] = {
        "mime": mime_type,
        "paragraphs": len(document.paragraphs),
        "tables": table_count,
    }
    return ParseOutcome(
        status="parsed",
        text=text,
        metadata=metadata,
        page_count=None,
        error=None if text else "Empty document",
    )


def _parse_csv(path: Path, mime_type: str) -> ParseOutcome:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error=str(exc),
        )
    metadata: dict[str, Any] = {"mime": mime_type}
    reader = csv.reader(io.StringIO(raw))
    rows: list[str] = []
    for idx, row in enumerate(reader):
        if idx >= CSV_PREVIEW_ROWS:
            rows.append("…")
            break
        rows.append(", ".join(cell.strip() for cell in row))
    text = _truncate_text("\n".join(rows))
    metadata["rows_sampled"] = min(len(rows), CSV_PREVIEW_ROWS)
    return ParseOutcome(
        status="parsed",
        text=text,
        metadata=metadata,
        page_count=None,
        error=None if text else "Empty CSV",
    )


def _parse_text(path: Path, mime_type: str) -> ParseOutcome:
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error=str(exc),
        )
    return ParseOutcome(
        status="parsed",
        text=_truncate_text(text),
        metadata={"mime": mime_type},
        page_count=None,
        error=None if text.strip() else "Empty text file",
    )


def _parse_xlsx(path: Path, mime_type: str) -> ParseOutcome:
    if openpyxl is None:
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error="openpyxl dependency is not installed",
        )
    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)  # type: ignore[attr-defined]
    except Exception as exc:  # noqa: BLE001
        return ParseOutcome(
            status="failed",
            text="",
            metadata={},
            page_count=None,
            error=str(exc),
        )
    lines: list[str] = []
    sheet_count = 0
    cell_count = 0
    for sheet in workbook.worksheets:  # type: ignore[attr-defined]
        sheet_count += 1
        lines.append(f"# Sheet: {sheet.title}")
        for row_index, row in enumerate(sheet.iter_rows(values_only=False), 1):  # type: ignore[attr-defined]
            if row_index > XLSX_PREVIEW_ROWS:
                lines.append(".")
                break
            cells = []
            for cell in row:
                value = cell.value
                if value is None:
                    continue
                text_value = str(value).strip()
                if not text_value:
                    continue
                cells.append(f"{sheet.title}!{cell.coordinate}: {text_value}")
                cell_count += 1
            if cells:
                lines.append("; ".join(cells))
        lines.append("")
    close_fn = getattr(workbook, "close", None)
    if callable(close_fn):
        close_fn()
    text_out = _truncate_text("\n".join(lines))
    metadata: dict[str, Any] = {
        "mime": mime_type,
        "sheets": sheet_count,
        "cells_sampled": cell_count,
    }
    return ParseOutcome(
        status="parsed",
        text=text_out,
        metadata=metadata,
        page_count=None,
        error=None if text_out.strip() else "Empty workbook",
    )


def _unsupported_cad(path: Path, mime_type: str) -> ParseOutcome:
    metadata: dict[str, Any] = {
        "mime": mime_type,
        "note": "Preview not available; metadata only",
        "file_name": path.name,
    }
    return ParseOutcome(
        status="parsed",
        text="",
        metadata=metadata,
        page_count=None,
        error=None,
    )
