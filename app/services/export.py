import csv
from pathlib import Path
from typing import List, Dict

try:
    import openpyxl  # type: ignore
except Exception:  # pragma: no cover
    openpyxl = None

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None


def export_csv(lines: List[Dict], path: Path) -> Path:
    with path.open("w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["desc", "qty", "unit", "unit_cost", "line_total"])
        for line in lines:
            writer.writerow(
                [
                    line.get("desc"),
                    line.get("qty"),
                    line.get("unit"),
                    line.get("unit_cost"),
                    line.get("line_total"),
                ]
            )
    return path


def export_excel(lines: List[Dict], path: Path) -> Path:
    if openpyxl is None:  # pragma: no cover - fallback
        return export_csv(lines, path.with_suffix(".csv"))
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["desc", "qty", "unit", "unit_cost", "line_total"])
    for line in lines:
        ws.append(
            [
                line.get("desc"),
                line.get("qty"),
                line.get("unit"),
                line.get("unit_cost"),
                line.get("line_total"),
            ]
        )
    wb.save(path)
    return path


def export_docx(summary: str, path: Path) -> Path:
    if Document is None:  # pragma: no cover - fallback
        path.write_text(summary)
        return path
    doc = Document()
    doc.add_heading("Estimate Summary", 0)
    doc.add_paragraph(summary)
    doc.save(path)
    return path
